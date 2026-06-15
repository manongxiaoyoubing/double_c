from docx import Document
from docx.shared import Pt, Inches, RGBColor
from db import db
from models import ReviewRecord, ReviewDetail
from datetime import datetime

class ReportGenerator:
    def generate_word(self, review_id, output_path):
        """生成Word格式报告"""
        record = ReviewRecord.query.get(review_id)
        details = ReviewDetail.query.filter_by(review_id=review_id).all()
        
        doc = Document()
        
        # 标题
        heading = doc.add_heading('双碳合规审查报告', 0)
        
        # 项目信息
        doc.add_heading('一、项目信息', 1)
        doc.add_paragraph(f'项目名称：{record.project_name}')
        doc.add_paragraph(f'项目类型：{record.project_type}')
        doc.add_paragraph(f'所属行业：{record.industry}')
        doc.add_paragraph(f'所属地区：{record.region}')
        doc.add_paragraph(f'提交时间：{record.submitted_at.strftime("%Y-%m-%d %H:%M") if record.submitted_at else ""}')
        
        # 审查结果
        doc.add_heading('二、审查结果', 1)
        doc.add_paragraph(f'审查结论：{record.review_result}')
        doc.add_paragraph(f'合规得分：{record.compliance_score}分')
        doc.add_paragraph(f'总体结论：{record.overall_conclusion}')
        
        # 详细检查项
        doc.add_heading('三、详细检查项', 1)
        for i, detail in enumerate(details, 1):
            p = doc.add_paragraph(f'{i}. {detail.check_item}')
            p.add_run(f'\n    检查结果：{"符合" if detail.is_compliant else "不符合"}')
            p.add_run(f'\n    结论：{detail.conclusion}')
        
        # 保存
        doc.save(output_path)
        return output_path
    
    def generate_pdf(self, review_id, output_path):
        """生成PDF格式报告（简化版，实际项目可用reportlab）"""
        # 简化实现：先生成Word再转换
        word_path = output_path.replace('.pdf', '.docx')
        self.generate_word(review_id, word_path)
        # 实际项目中可调用libreoffice或docx2pdf进行转换
        return word_path
